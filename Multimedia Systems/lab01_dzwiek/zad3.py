import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import soundfile as sf
from docx import Document
from docx.shared import Inches
from io import BytesIO

def plot(data,fs, fsize,margins, axs):
    time = np.arange(data.shape[0]) / fs
    axs[0].plot(time, data)
    axs[0].set_title('Sygnał w dziedzinie czasu')
    axs[0].set_xlabel('Czas [s]')
    axs[0].set_xlim(margins)
    axs[0].set_ylabel('Amplituda')
    axs[0].grid(True)

    yf = scipy.fftpack.fft(data, fsize)
    
    freqs = np.arange(0, fs / 2, fs / fsize)
    amplitudes = np.abs(yf[:fsize//2])
    amplitudes_db = 20 * np.log10(amplitudes + 1e-10)
    
    axs[1].plot(freqs, amplitudes_db)
    axs[1].set_title(f'Widmo amplitudowe (fsize={fsize})')
    axs[1].set_xlabel('Częstotliwość [Hz]')
    axs[1].set_xlim(margins)
    axs[1].set_ylabel('Amplituda [dB]')
    axs[1].grid(True)

    max_idx = np.argmax(amplitudes_db)
    max_freq = freqs[max_idx]
    max_amp = amplitudes_db[max_idx]

    return max_freq, max_amp


def generate_report():
    document = Document()
    document.add_heading('Analiza sygnałów dźwiękowych - Zadanie 3', 0)

    files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav', 'sin_combined.wav']
    
    fsizes = [2**8, 2**12, 2**16]
    margins = [[0,0.02],[0.133,0.155]]

    for file in files:
        document.add_heading(f'Plik: {file}', 2)
        data, fs = sf.read(file, dtype=np.int32)
        
        for fsize in fsizes:
            for margin in margins:
                document.add_heading('Time margin {}'.format(margin),3)
                #document.add_heading(f'fsize = {fsize}', 3)
                
                fig, axs = plt.subplots(2, 1, figsize=(10, 7))
                max_freq, max_amp = plot(data,fs, fsize,margin,axs)
                
                fig.suptitle(f'Analiza pliku {file} dla fsize={fsize}')
                fig.tight_layout(pad=1.5)
                memfile = BytesIO()
                fig.savefig(memfile, format='png')
                memfile.seek(0)
                document.add_picture(memfile, width=Inches(6))
                memfile.close()
                plt.close(fig)
                document.add_paragraph(f'Wartość maksymalnej amplitudy: {max_amp:.2f} dB')
                document.add_paragraph(f'Częstotliwość prążka dla najwyższej amplitudy: {max_freq:.2f} Hz')

    document.save('report.docx')

if __name__ == '__main__':
    generate_report()