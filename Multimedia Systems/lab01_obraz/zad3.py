import pandas as pd
import numpy as np
import matplotlib.pylab as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO

df = pd.DataFrame(data={
    'Filename': ['B01.png'],
    'Grayscale': [False],
    'Fragments': [
        [[50, 50, 250, 250], [100, 200, 300, 400]]
    ]
})

def show_plot_3x3(img_fragment):
    R = img_fragment[:,:,0]
    G = img_fragment[:,:,1]
    B = img_fragment[:,:,2]

    Y1 = 0.299 * R + 0.587 * G + 0.114 * B 
    Y2 = 0.2126 * R + 0.7152 * G + 0.0722 * B

    img_R_color = img_fragment.copy()
    img_R_color[:,:,1] = 0
    img_R_color[:,:,2] = 0

    img_G_color = img_fragment.copy()
    img_G_color[:,:,0] = 0
    img_G_color[:,:,2] = 0

    img_B_color = img_fragment.copy()
    img_B_color[:,:,0] = 0
    img_B_color[:,:,1] = 0

    if np.issubdtype(img_fragment.dtype, np.integer): 
        v_min, v_max = 0, 255
    else:
        v_min, v_max = 0.0, 1.0

    fig, axs = plt.subplots(3, 3, figsize=(10, 10))

    axs[0,0].imshow(img_fragment); axs[0,0].set_title('O')
    axs[0,1].imshow(Y1, cmap='gray', vmin=v_min, vmax=v_max); axs[0,1].set_title('Y1')
    axs[0,2].imshow(Y2, cmap='gray', vmin=v_min, vmax=v_max); axs[0,2].set_title('Y2')

    axs[1,0].imshow(R, cmap='gray', vmin=v_min, vmax=v_max); axs[1,0].set_title('R (Gray)')
    axs[1,1].imshow(G, cmap='gray', vmin=v_min, vmax=v_max); axs[1,1].set_title('G (Gray)')
    axs[1,2].imshow(B, cmap='gray', vmin=v_min, vmax=v_max); axs[1,2].set_title('B (Gray)')

    axs[2,0].imshow(img_R_color); axs[2,0].set_title('R (Color)')
    axs[2,1].imshow(img_G_color); axs[2,1].set_title('G (Color)')
    axs[2,2].imshow(img_B_color); axs[2,2].set_title('B (Color)')

    return fig

def generate_report(df):
    document = Document()
    document.add_heading('Raport', 0)

    for index, row in df.iterrows():
        filename = row['Filename']
        document.add_heading(f'Plik: {filename}', 2)
        
        img = plt.imread(filename)
            
        if row['Fragments'] is not None:
            for i, f in enumerate(row['Fragments']):
                fragment = img[f[0]:f[2], f[1]:f[3]].copy()
                
                document.add_heading(f'Fragment {i+1} (Współrzędne: {f})', 3)
                
                fig = show_plot_3x3(fragment)
                
                fig.suptitle(f'Analiza fragmentu {f} z pliku {filename}', fontsize=14)
                fig.tight_layout(pad=1.5)
                
                memfile = BytesIO()
                fig.savefig(memfile, format='png')
                memfile.seek(0)
                
                document.add_picture(memfile, width=Inches(6))
                
                memfile.close()
                plt.close(fig)

    report_name = 'report.docx'
    document.save(report_name)
    print(f"Raport: {report_name}")

generate_report(df)
