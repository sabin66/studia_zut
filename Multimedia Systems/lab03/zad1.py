import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import cv2
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO

##########################################
### Settings #############################
##########################################

Test=True
ColorFit_Test=True


GrayScale_bits = [1,2,4]


pallet8 = np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 0.0, 1.0,],
        [0.0, 1.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [1.0, 0.0, 0.0,],
        [1.0, 0.0, 1.0,],
        [1.0, 1.0, 0.0,],
        [1.0, 1.0, 1.0,],
])
pallet16 =  np.array([
        [0.0, 0.0, 0.0,], 
        [0.0, 1.0, 1.0,],
        [0.0, 0.0, 1.0,],
        [1.0, 0.0, 1.0,],
        [0.0, 0.5, 0.0,], 
        [0.5, 0.5, 0.5,],
        [0.0, 1.0, 0.0,],
        [0.5, 0.0, 0.0,],
        [0.0, 0.0, 0.5,],
        [0.5, 0.5, 0.0,],
        [0.5, 0.0, 0.5,],
        [1.0, 0.0, 0.0,],
        [0.75, 0.75, 0.75,],
        [0.0, 0.5, 0.5,],
        [1.0, 1.0, 1.0,], 
        [1.0, 1.0, 0.0,]
])

Color_pallets=[
    pallet8,pallet16
]

M2=np.array([[0,8,2,10],
             [12,4,14,6],
             [3,11,1,9],
             [12,7,13,5]])

OutputRaportFile = ".docx" 

##########################################
### Data Set #############################
##########################################

ImgDir = r'.' # Address of folder with files (do nor delete `r``)


GsImages=["IMG_GS/GS_0002.png"] # list of file names with GS images
ColorImages=["IMG_SMALL/SMALL_0003.png"] # list of file names of Color images

##########################################
### Functions to  ########################
##########################################

def imgToUint8(img):
    if not np.issubdtype(img.dtype,np.unsignedinteger):
        img = (img*255).astype('uint8')
    return img

def imgToFloat(img):
    if not np.issubdtype(img.dtype,np.floating):
        img = img/255.0
    return img

def colorFit(pixel,Pallet):
        v = Pallet-pixel
        euc_dist = np.linalg.norm(v,axis=1)
        idx = np.argmin(euc_dist)
        return Pallet[idx]

def kwant_colorFit(img,Pallet):
        out_img = img.copy()
        for k in range(img.shape[1]):
                for w in range(img.shape[0]):
                        tmp = colorFit(img[w,k],Pallet)
                        if len(tmp)==1:
                            out_img[w,k]=tmp[0]
                        else:
                            out_img[w,k]=tmp[:]
        return out_img.astype(img.dtype)

def dith_randm(img):
        out_img = img.copy()
        r = np.random.rand(img.shape[0],img.shape[1])
        out_img = (out_img>=r).astype(img.dtype)
        out_img *= 1
              
        return out_img.astype(img.dtype)

def dith_ordered(img,Pallet,r=1,M=M2):
        out_img = img.copy()
        ##### 
        return out_img.astype(img.dtype)

def dith_FS(img,Pallet):
        out_img = img.copy()
        ##### 
        return out_img.astype(img.dtype)



##########################################
### Main Program  ########################
##########################################

def process_and_plot_GS(img,bit,filename,counter,figsize=(5,5)):
        if len(img.shape)>2:
                img=img[:,:,0]
        palett=np.linspace(0,1,2**bit).reshape(-1,1)
        qwant_img=kwant_colorFit(img,palett)

        order_img=dith_ordered(img,palett)
        FS_img=dith_FS(img,palett)
        if bit==1:
            rand_img=dith_randm(img)
            f,axs=plt.subplots(2,3,num=counter,figsize=figsize) 
            f.suptitle(f"{filename} Dithering 1-bit")
            axs[0,0].imshow(img,cmap="gray")
            axs[0,0].set_title("Oryginał")
            axs[0,0].set_axis_off()

            axs[1,0].remove()

            axs[0,2].imshow(rand_img,cmap="gray")
            axs[0,2].set_title("Dithering\n Losowy")
            axs[0,2].set_axis_off()

            axt=[axs[0,1],axs[1,1],axs[1,2]]
        else:
            f,axs=plt.subplots(1,4,num=counter,figsize=figsize) 
            f.suptitle(f"{filename} Dithering {bit}-bitów")
            axs[0].imshow(img,cmap="gray")
            axs[0].set_title("Oryginał")
            axs[0].set_axis_off()  
            axt=[axs[1],axs[2],axs[3]] 
            rand_img=0

        axt[0].imshow(qwant_img,cmap="gray")
        axt[0].set_title("Kwantyzacja")
        axt[0].set_axis_off()

        
        axt[1].imshow(order_img,cmap="gray")
        axt[1].set_title("Dithering\n Zorganizowany")
        axt[1].set_axis_off()

        
        axt[2].imshow(FS_img,cmap="gray")
        axt[2].set_title("Dithering\n Floyda-Steinberga")
        axt[2].set_axis_off()

        if Test:
                if bit==1:
                    print(f"Test of uniqe values counts:\n"+
                        f"Unique values in Pallet: {np.unique(palett).size}\n"+
                        f"Dithering Random: {np.unique(rand_img).size}\n"+
                        f"Dithering Ordered: {np.unique(order_img).size}\n"+
                        f"Dithering Floyd-Steinberg: {np.unique(FS_img).size}\n")
                else:
                    print(f"Test of uniqe values counts:\n"+
                        f"Unique values in Pallet: {np.unique(palett).size}\n"+
                        f"Dithering Ordered: {np.unique(order_img).size}\n"+
                        f"Dithering Floyd-Steinberg: {np.unique(FS_img).size}\n")
        return f
                    

def process_and_plot_Color(img,palett,filename,counter,figsize=(5,5)):
        if img.shape[2]>3:
                img=img[:,:,:3]
        qwant_img=kwant_colorFit(img,palett)

        order_img=dith_ordered(img,palett)
        FS_img=dith_FS(img,palett)

        f,axs=plt.subplots(1,4,num=counter,figsize=figsize) 
        f.suptitle(f"{filename} Dithering {len(palett)} kolorów")
        axs[0].imshow(img)
        axs[0].set_title("Oryginał")
        axs[0].set_axis_off()  

        axs[1].imshow(qwant_img)
        axs[1].set_title("Kwantyzacja")
        axs[1].set_axis_off()  

        axs[2].imshow(order_img)
        axs[2].set_title("Dithering\n Zorganizowany")
        axs[2].set_axis_off()  

        axs[3].imshow(FS_img)
        axs[3].set_title("Dithering\n Floyda-Steinberga")
        axs[3].set_axis_off() 

        return f




if Test:
        if ColorFit_Test:
            paleta = np.linspace(0,1,3).reshape(3,1)
            print(f"0.43 -> {colorFit(0.43,paleta)}") 
            print(f"0.66 -> {colorFit(0.66,paleta)}") 
            print(f"0.8 -> {colorFit(0.8,paleta)}") 

            print(f"[0.25,0.25,0.5] 8 kolorów -> {colorFit(np.array([0.25,0.25,0.5]),pallet8)}")
            print(f"[0.25,0.25,0.5] 16 kolorów-> {colorFit(np.array([0.25,0.25,0.5]),pallet16)}")

              
        file=GsImages[0]
        img=imgToFloat(plt.imread(os.path.join(ImgDir,file)))
        counter=1
        for bit in GrayScale_bits:
               process_and_plot_GS(img=img,bit=bit,filename=file,counter=counter)
               counter+=1

        file=ColorImages[0]
        img=imgToFloat(plt.imread(os.path.join(ImgDir,file)))
        for palett in Color_pallets:
               process_and_plot_Color(img=img,palett=palett,filename=file,counter=counter)
               counter+=1
        
        plt.show()
else:
    # generate raport
    document = Document()
    document.add_heading('Report',0) # tworzenie nagłówków druga wartość to poziom nagłówka 
    document.add_paragraph("Autor: ")
    document.add_paragraph("Proszę wstawić mi 2 jeżeli tego nie wyedytuję")
    document.add_section()
    document.add_heading("Test ditheringu na obrazach w skali odcieni szarości",1)
    counter = 1 
    for file in GsImages:
        img=imgToFloat(plt.imread(os.path.join(ImgDir,file)))
        for bit in GrayScale_bits:
                f = process_and_plot_GS(img=img,bit=bit,filename=file,counter=counter)
                memfile = BytesIO() 
                f.savefig(memfile)
                document.add_picture(memfile, width=Inches(6)) # set document size
                memfile.close()
                f.clf()
    document.add_section()
    document.add_heading("Test ditheringu na obrazach kolorowych",1)
    for file in ColorImages:
        img=imgToFloat(plt.imread(os.path.join(ImgDir,file)))
        for palett in Color_pallets:
            f = process_and_plot_Color(img=img,palett=palett,filename=file,counter=counter)
            memfile = BytesIO() 
            f.savefig(memfile)
            document.add_picture(memfile, width=Inches(6)) # set document size
            memfile.close()
            f.clf()
    document.add_section()
    document.add_heading("Podsumowanie i wnioski",1)
    document.add_paragraph("Tu proszę zebrać wszystkie obserwacje na podstawie powyższych wykresów i napisać wnioski.")
    document.save(OutputRaportFile) 