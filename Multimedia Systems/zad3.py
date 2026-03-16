import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import soundfile as sf
from docx import Document
from docx.shared import Inches
from io import BytesIO

def analyze_and_plot(filename, fsize, axs):
    data, fs = sf.read(filename, dtype=np.int32)
    
    if len(data.shape) > 1:
        data = data[:, 0]

    time = np.arange(data.shape[0]) / fs
    axs[0].plot(time, data)
    axs[0].set_title('Sygnał w dziedzinie czasu')
    axs[0].set_xlabel('Czas [s]')
    axs[0].set_ylabel('Amplituda')
    axs[0].grid(True)

    yf = scipy.fftpack.fft(data, fsize)
    
    freqs = np.arange(0, fs / 2, fs / fsize)
    
    amplitudes = np.abs(yf[:fsize//2])
    amplitudes_db = 20 * np.log10(amplitudes + 1e-10)
    
    axs[1].plot(freqs, amplitudes_db)
    axs[1].set_title(f'Widmo amplitudowe (fsize={fsize})')
    axs[1].set_xlabel('Częstotliwość [Hz]')
    axs[1].set_ylabel('Amplituda [dB]')
    axs[1].grid(True)

    max_idx = np.argmax(amplitudes_db)
    max_freq = freqs[max_idx]
    max_amp = amplitudes_db[max_idx]   # Wartość amplitudy w dB

    return max_freq, max_amp


def generate_report():
    document = Document()
    document.add_heading('Analiza sygnałów dźwiękowych - Zadanie 3', 0)

    # Dokładne nazwy plików z rozszerzeniem i podkreśleniami zgodnie z listą
    files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav', 'sin_combined.wav']
    
    # 3 różne rozmiary okna FFT z instrukcji
    fsizes = [2**8, 2**12, 2**16]

    for file in files:
        document.add_heading(f'Plik: {file}', 2)
        
        for fsize in fsizes:
            document.add_heading(f'fsize = {fsize}', 3)
            
            # Tworzenie plota z dwiema osiami
            fig, axs = plt.subplots(2, 1, figsize=(10, 7))
            
            # Analiza i rysowanie (funkcja modyfikuje obiekt axs)
            max_freq, max_amp = analyze_and_plot(file, fsize, axs)
            
            # Poprawa czytelności wykresów
            fig.suptitle(f'Analiza pliku {file} dla fsize={fsize}')
            fig.tight_layout(pad=1.5)
            
            # Zapisywanie figury do bufora w pamięci
            memfile = BytesIO()
            fig.savefig(memfile, format='png')
            memfile.seek(0) # Przesunięcie "wskaźnika" na początek pliku w pamięci
            
            # Dodanie obrazu do dokumentu Word
            document.add_picture(memfile, width=Inches(6))
            memfile.close()
            plt.close(fig) # Zamknięcie obiektu figury zwalnia pamięć RAM
            
            # Automatyczne dodanie do raportu tekstowych wyników analizy widmowej
            document.add_paragraph(f'Wartość maksymalnej amplitudy: {max_amp:.2f} dB')
            document.add_paragraph(f'Częstotliwość prążka dla najwyższej amplitudy: {max_freq:.2f} Hz')

    # Zapis raportu do formatu .docx
    report_filename = 'report.docx'
    document.save(report_filename)
    print(f'Wygenerowano plik {report_filename}.')

if __name__ == '__main__':
    generate_report()