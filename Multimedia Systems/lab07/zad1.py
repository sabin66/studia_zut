import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import cv2
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO
import scipy.fftpack

##########################################
### Settings #############################
##########################################

Test=True

OutputRaportFile = "yes.docx" 

Chroma_options=["4:4:4","4:2:2"]
Quant_options=[True,False]

QY= np.array([
        [16, 11, 10, 16, 24,  40,  51,  61],
        [12, 12, 14, 19, 26,  58,  60,  55],
        [14, 13, 16, 24, 40,  57,  69,  56],
        [14, 17, 22, 29, 51,  87,  80,  62],
        [18, 22, 37, 56, 68,  109, 103, 77],
        [24, 36, 55, 64, 81,  104, 113, 92],
        [49, 64, 78, 87, 103, 121, 120, 101],
        [72, 92, 95, 98, 112, 100, 103, 99],
        ])

QC= np.array([
        [17, 18, 24, 47, 99, 99, 99, 99],
        [18, 21, 26, 66, 99, 99, 99, 99],
        [24, 26, 56, 99, 99, 99, 99, 99],
        [47, 66, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        ])

QN= np.ones((8,8))

##########################################
### Data Set #############################
##########################################

ImgDir = r'.' # Address of folder with files (do nor delete `r``)

Images=[ #list of dictionaries
    {
        "Filename":"webcam-toy-photo22.jpg", # File name
        "ROIs":[[250,500,128,128]] # list of Region of interests for this image more then 1 per file
    }
]


##########################################
### Functions to  ########################
##########################################

def dct2(a):
    return scipy.fftpack.dct( scipy.fftpack.dct( a.astype(float), axis=0, norm='ortho' ), axis=1, norm='ortho' )

def idct2(a):
    return scipy.fftpack.idct( scipy.fftpack.idct( a.astype(float), axis=0 , norm='ortho'), axis=1 , norm='ortho')

def zigzag(A):
    template= np.array([
            [0,  1,  5,  6,  14, 15, 27, 28],
            [2,  4,  7,  13, 16, 26, 29, 42],
            [3,  8,  12, 17, 25, 30, 41, 43],
            [9,  11, 18, 24, 31, 40, 44, 53],
            [10, 19, 23, 32, 39, 45, 52, 54],
            [20, 22, 33, 38, 46, 51, 55, 60],
            [21, 34, 37, 47, 50, 56, 59, 61],
            [35, 36, 48, 49, 57, 58, 62, 63],
            ])
    if len(A.shape)==1:
        B=np.zeros((8,8))
        for r in range(0,8):
            for c in range(0,8):
                B[r,c]=A[template[r,c]]
    else:
        B=np.zeros((64,))
        for r in range(0,8):
            for c in range(0,8):
                B[template[r,c]]=A[r,c]
    return B

def rle_encode(data):
    encoded = []
    i = 0
    while i < len(data):
        val = data[i]
        count = 1
        while i + count < len(data) and data[i + count] == val and count < 127:
            count += 1
        encoded.append(count)
        encoded.append(val)
        i += count
    return np.array(encoded)

def rle_decode(encoded):
    decoded = []
    i = 0
    while i < len(encoded):
        count = int(encoded[i])
        val = encoded[i + 1]
        decoded.extend([val] * count)
        i += 2
    return np.array(decoded)

def chroma_subsample(layer, ratio):
    if ratio == "4:2:2":
        # co druga kolumna
        return layer[:, ::2]
    else:  # "4:4:4"
        return layer.copy()

def chroma_upsample(layer, ratio):
    if ratio == "4:2:2":
        return np.repeat(layer, 2, axis=1)
    else:  # "4:4:4"
        return layer.copy()

# JPEG container definition
class JPEG_class:
    def __init__(self,Y,Cb,Cr,OGShape,Ratio="4:4:4",QY=np.ones((8,8)),QC=np.ones((8,8))):
        self.shape = OGShape
        self.Y=Y
        self.Cb=Cb
        self.Cr=Cr
        self.ChromaRatio=Ratio
        self.QY=QY
        self.QC=QC

def CompressBlock(block, Q):
    block = block - 128
    d = dct2(block)
    qd = np.round(d / Q).astype(int)
    vector = zigzag(qd)
    return vector

def DecompressBlock(vector, Q):
    qd = zigzag(vector)
    d = qd * Q
    block = idct2(d)
    block = block + 128
    return block

def CompressLayer(L, Q):
    S=np.array([])
    for w in range(0, L.shape[0], 8):
        for k in range(0, L.shape[1], 8):
            block = L[w:(w+8), k:(k+8)]
            S = np.append(S, CompressBlock(block, Q))
    return S

def DecompressLayer(S, Q, layer_shape):
    L = np.zeros(layer_shape)
    for idx, i in enumerate(range(0, S.shape[0], 64)):
        vector = S[i:(i+64)]
        m = L.shape[1] / 8
        k = int((idx % m) * 8)
        w = int((idx // m) * 8)
        L[w:(w+8), k:(k+8)] = DecompressBlock(vector, Q)
    return L

def CompressJPEG(RGB, Ratio="4:4:4", QY=QN, QC=QN):
    YCrCb = cv2.cvtColor(RGB, cv2.COLOR_RGB2YCrCb).astype(int)
    
    y_layer = YCrCb[:,:,0].astype(float)
    cr_layer = YCrCb[:,:,1].astype(float)
    cb_layer = YCrCb[:,:,2].astype(float)

    cr_sub = chroma_subsample(cr_layer, Ratio)
    cb_sub = chroma_subsample(cb_layer, Ratio)

    Y_compressed = CompressLayer(y_layer, QY)
    Cr_compressed = CompressLayer(cr_sub, QC)
    Cb_compressed = CompressLayer(cb_sub, QC)
    
    Y_rle = rle_encode(Y_compressed)
    Cr_rle = rle_encode(Cr_compressed)
    Cb_rle = rle_encode(Cb_compressed)

    JPEG = JPEG_class(Y_rle, Cb_rle, Cr_rle, RGB.shape, Ratio=Ratio, QY=QY, QC=QC)
    JPEG.YShape = y_layer.shape
    JPEG.CrShape = cr_sub.shape
    JPEG.CbShape = cb_sub.shape

    JPEG.Y_raw_len = len(Y_compressed)
    JPEG.Cr_raw_len = len(Cr_compressed)
    JPEG.Cb_raw_len = len(Cb_compressed)
    JPEG.Y_rle_len = len(Y_rle)
    JPEG.Cr_rle_len = len(Cr_rle)
    JPEG.Cb_rle_len = len(Cb_rle)

    return JPEG
    
def DecompressJPEG(JPEG):
    Y_compressed = rle_decode(JPEG.Y)
    Cr_compressed = rle_decode(JPEG.Cr)
    Cb_compressed = rle_decode(JPEG.Cb)
    
    Y = DecompressLayer(Y_compressed, JPEG.QY, JPEG.YShape)
    Cr = DecompressLayer(Cr_compressed, JPEG.QC, JPEG.CrShape)
    Cb = DecompressLayer(Cb_compressed, JPEG.QC, JPEG.CbShape)
    
    Cr = chroma_upsample(Cr, JPEG.ChromaRatio)
    Cb = chroma_upsample(Cb, JPEG.ChromaRatio)
    
    YCrCb = np.dstack([Y, Cr, Cb]).clip(0, 255).astype(np.uint8)
    
    RGB = cv2.cvtColor(YCrCb, cv2.COLOR_YCrCb2RGB)
    
    return RGB

##########################################
### Main Program  ########################
##########################################

def plot_comparisone(counter, OG,Decomp,figsize=(5,8) ):
    fig, axs = plt.subplots(4, 2 ,num=counter, sharex=True, sharey=True,figsize=figsize )
    # obraz oryginalny 
    axs[0,0].imshow(OG) #RGB 
    PRZED_YCrCb=cv2.cvtColor(OG,cv2.COLOR_RGB2YCrCb)
    axs[1,0].imshow(PRZED_YCrCb[:,:,0],cmap='gray') 
    axs[2,0].imshow(PRZED_YCrCb[:,:,1],cmap='gray')
    axs[3,0].imshow(PRZED_YCrCb[:,:,2],cmap='gray')
    
    axs[0,0].set_title("Oryginał")
    axs[1,0].set_title("Y")
    axs[2,0].set_title("Cr")
    axs[3,0].set_title("Cb")

    # obraz po dekompresji
    axs[0,1].imshow(Decomp) #RGB 
    PO_YCrCb=cv2.cvtColor(Decomp,cv2.COLOR_RGB2YCrCb)
    axs[1,1].imshow(PO_YCrCb[:,:,0],cmap='gray')
    axs[2,1].imshow(PO_YCrCb[:,:,1],cmap='gray')
    axs[3,1].imshow(PO_YCrCb[:,:,2],cmap='gray')
    
    axs[0,1].set_title("Po kompresji")
    axs[1,1].set_title("Y")
    axs[2,1].set_title("Cr")
    axs[3,1].set_title("Cb")
    
    for ax in axs.flatten():
        ax.set_axis_off() 
    
    return fig

if Test:
    img=plt.imread(os.path.join(ImgDir,Images[0]["Filename"]))
    ROI=Images[0]["ROIs"][0]
    fragment= img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2]]
    Counter=1
    for Chroma in Chroma_options:
        for Quant in Quant_options:
            if Quant:
                tJPEG=CompressJPEG(fragment,Ratio=Chroma,QY=QY,QC=QC)
            else:
                tJPEG=CompressJPEG(fragment,Ratio=Chroma,QY=QN,QC=QN)
                
            # Wypisanie informacji o kompresji RLE
            print(f"=== Chroma: {Chroma}, Kwantyzacja: {Quant} ===")
            print(f"  Y:  przed RLE={tJPEG.Y_raw_len}, po RLE={tJPEG.Y_rle_len}, redukcja={100*(1-tJPEG.Y_rle_len/tJPEG.Y_raw_len):.1f}%")
            print(f"  Cr: przed RLE={tJPEG.Cr_raw_len}, po RLE={tJPEG.Cr_rle_len}, redukcja={100*(1-tJPEG.Cr_rle_len/tJPEG.Cr_raw_len):.1f}%")
            print(f"  Cb: przed RLE={tJPEG.Cb_raw_len}, po RLE={tJPEG.Cb_rle_len}, redukcja={100*(1-tJPEG.Cb_rle_len/tJPEG.Cb_raw_len):.1f}%")
            
            New_Fragment=DecompressJPEG(tJPEG)
                
            f = plot_comparisone(Counter,fragment,New_Fragment)
            f.suptitle("Plik: {} Chroma: {} Kwantyzacja: {}".format(Images[0]["Filename"],Chroma,Quant))
            Counter+=1
    plt.show()
    
else:
    # generate raport
    document = Document()
    document.add_heading('Report',0) # tworzenie nagłówków druga wartość to poziom nagłówka 
    document.add_paragraph("Autor: ")
    document.add_section()
    document.add_heading("Fragmenty wygenerowane na podstawie działania funkcji",1)
    Counter=1
    for file_dict in Images:
        filename=file_dict['Filename']
        img=plt.imread(os.path.join(ImgDir,filename))
        for ROI in file_dict['ROIs']:
            fragment= img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2]]
            
            for Chroma in Chroma_options:
                for Quant in Quant_options:
                    if Quant:
                        tJPEG=CompressJPEG(fragment,Ratio=Chroma,QY=QY,QC=QC)
                    else:
                        tJPEG=CompressJPEG(fragment,Ratio=Chroma,QY=QN,QC=QN)
                        
                    New_Fragment=DecompressJPEG(tJPEG)
                        
                    f = plot_comparisone(Counter,fragment,New_Fragment)
                    f.suptitle("Plik: {} Chroma: {} Kwantyzacja: {}".format(filename,Chroma,Quant))
                    memfile = BytesIO() 
                    f.savefig(memfile)
                    document.add_picture(memfile, width=Inches(6)) # set document size
                    
                    # Dodanie informacji o kompresji RLE
                    document.add_paragraph(
                        f"Chroma: {Chroma}, Kwantyzacja: {Quant}\n"
                        f"  Y:  przed RLE={tJPEG.Y_raw_len}, po RLE={tJPEG.Y_rle_len}, redukcja={100*(1-tJPEG.Y_rle_len/tJPEG.Y_raw_len):.1f}%\n"
                        f"  Cr: przed RLE={tJPEG.Cr_raw_len}, po RLE={tJPEG.Cr_rle_len}, redukcja={100*(1-tJPEG.Cr_rle_len/tJPEG.Cr_raw_len):.1f}%\n"
                        f"  Cb: przed RLE={tJPEG.Cb_raw_len}, po RLE={tJPEG.Cb_rle_len}, redukcja={100*(1-tJPEG.Cb_rle_len/tJPEG.Cb_raw_len):.1f}%"
                    )
                    
                    memfile.close()
                    f.clf()
                    Counter+=1
    document.add_section()
    document.add_heading("Podsumowanie i wnioski",1)
    document.add_paragraph("Tu proszę zebrać wszystkie obserwacje na podstawie powyższych wykresów i napisać wnioski.")
    document.save(OutputRaportFile) 