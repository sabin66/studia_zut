import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import cv2
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO
import scipy
import soundfile as sf
from scipy.interpolate import interp1d

##########################################
### Settings #############################
##########################################

Test = False
Kwant_Test = False # test Kwant function

Interpolation_kind = ["linear", "cubic"]

PlotSettings = {
    "Bits": [4, 8, 16, 24],
    "Decimation": [2, 4, 6, 10, 24],
    "InterpolationFrequency": [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]
}
ListeningSettings = {
    "Bits": [4, 8],
    "Decimation": [4, 6, 10, 24],
    "InterpolationFrequency": [4000, 8000, 11999, 16000, 16953]
}

OutputRaportFile = "Raport_Lab05.docx" 
OutputFolder = "."

##########################################
### Data Set #############################
##########################################

AudioDir = r'.'

SinFiles = [
    {"File": "SIN/sin_60Hz.wav", "TimeMargin": [0, 0.02]}, # +
    {"File": "SIN/sin_440Hz.wav", "TimeMargin": [0, 0.02]}, # ok
    {"File": "SIN/sin_8000Hz.wav", "TimeMargin": [0, 0.02]}, # -
    {"File": "SIN/sin_combined.wav", "TimeMargin": [0, 0.02]} # ok
] 
SingFiles = ["SING/sing_high1.wav","SING/sing_low1.wav","SING/sing_medium1.wav"] 

##########################################
### Functions ############################
##########################################

def plotAudio(Signal, Fs, axs, fsize, TimeMargin):
    time = np.arange(Signal.shape[0]) / Fs
    axs[0].plot(time, Signal)
    axs[0].set_title('Sygnał w dziedzinie czasu')
    axs[0].set_xlabel('Czas [s]')
    axs[0].set_xlim(TimeMargin)
    axs[0].set_ylabel('Amplituda')
    axs[0].grid(True)

    N = min(fsize, len(Signal))
    if N > 0:
        yf = scipy.fftpack.fft(Signal, fsize)
        
        freqs = np.arange(0, Fs / 2, Fs / fsize)
        amplitudes = np.abs(yf[:fsize//2])
        amplitudes_db = 20 * np.log10(amplitudes + 1e-10)
        
        axs[1].plot(freqs, amplitudes_db)
        axs[1].set_title(f'Widmo amplitudowe (fsize={fsize})')
        axs[1].set_xlabel('Częstotliwość [Hz]')
        axs[1].set_ylabel('Amplituda [dB]')
        axs[1].grid(True)

        max_idx = np.argmax(amplitudes_db)
        max_freq = freqs[max_idx]
        max_amp = amplitudes_db[max_idx]
        
        min_idx = np.argmin(amplitudes_db)
        min_freq = freqs[min_idx]
        min_amp = amplitudes_db[min_idx]
        # dzieki gemini za dopiski na wykresie
        axs[1].plot(max_freq, max_amp, 'ro') 
        axs[1].annotate(f'Max: {max_amp:.1f} dB ({max_freq:.1f} Hz)', 
                        xy=(max_freq, max_amp), xytext=(5, 5),
                        textcoords='offset points', color='red')
                        
        axs[1].plot(min_freq, min_amp, 'bo') 
        axs[1].annotate(f'Min: {min_amp:.1f} dB ({min_freq:.1f} Hz)', 
                        xy=(min_freq, min_amp), xytext=(5, -15),
                        textcoords='offset points', color='blue')

    return axs, max_freq, max_amp, min_freq, min_amp

def Kwant(data, bit):
    d = (2 ** bit) - 1 
    
    if np.issubdtype(data.dtype, np.floating):
        d_min = np.min(data)
        d_max = np.max(data)
    else:
        d_min = np.iinfo(data.dtype).min
        d_max = np.iinfo(data.dtype).max

    DataF = data.astype(float)

    if d_max > d_min:
        DataF = (DataF - d_min) / (d_max - d_min)
        DataF = np.round(DataF * d) / d
        DataF = DataF * (d_max - d_min) + d_min
    return DataF.astype(data.dtype)

def decimation(Signal, Fs, step):
    NewSignal = Signal[::step]
    NewFs = Fs // step
    return NewSignal, NewFs

def interpolation(Signal, Fs, NewFs, kind):
    if kind == "": 
        kind = "linear"
        
    duration = len(Signal) / Fs
    
    t_old = np.linspace(0, duration, len(Signal), endpoint=False)
    num_samples_new = int(len(Signal) * (NewFs / Fs))
    t_new = np.linspace(0, duration, num_samples_new, endpoint=False)
    
    f_interp = interp1d(t_old, Signal, kind=kind, fill_value="extrapolate")
    NewSignal = f_interp(t_new)
    
    return NewSignal.astype(Signal.dtype)


##########################################
### Main Program  ########################
##########################################

if Test:
    counter=1
    if Kwant_Test:
        T_X=[
            np.round(np.linspace(0,255,255,dtype=np.uint8)),
            np.round(np.linspace(np.iinfo(np.int32).min,np.iinfo(np.int32).max,1000,dtype=np.int32)),
            np.linspace(-1,1,10000),
        ]
        Bits=[1,2,4]
        for X in T_X:
            for bit in Bits:
                kwanted=Kwant(X,bit)
                print(f"Bits {bit} == {2**bit} values, unique values {np.unique(kwanted).size}. Dtype before {X.dtype} and after {kwanted.dtype}")
                plt.figure(counter)
                plt.plot(X,kwanted)
                plt.title(f"{bit} bit - typ {X.dtype}")
                counter+=1
                
    else:
        file=SinFiles[0]
        Signal, Fs = sf.read(os.path.join(AudioDir,file["File"]), dtype='float32') 
        # test decimation
        dec_Signal,dec_Fs=decimation(Signal,Fs,10)
        f,axs=plt.subplots(2,1,num=counter,figsize=(5,5)) 
        counter+=1
        plotAudio(Signal=dec_Signal,Fs=dec_Fs,axs=axs,fsize=2**12,TimeMargin=file["TimeMargin"])
        f.suptitle(f"{file['File']} Decymacja, krok 10")
        
        # test interpolation
        for kind in Interpolation_kind:
            Int_Fs=16000
            Int_Signal=interpolation(Signal=Signal,Fs=Fs,NewFs=Int_Fs,kind=kind)
            f,axs=plt.subplots(2,1,num=counter,figsize=(5,5)) 
            counter+=1
            plotAudio(Signal=Int_Signal,Fs=Int_Fs,axs=axs,fsize=2**12,TimeMargin=file["TimeMargin"])
            f.suptitle(f"{file['File']} Interpolacja {kind}")
        
    plt.show()
    
else:
    # generate raport
    document = Document()
    document.add_heading('Report',0)
    document.add_paragraph("Autor: Dorian Sobierański (uwaga - wnioski na samym dole grupowo)")
    document.add_section()
    document.add_heading("Sprawdzanie działania napisanych funkcji na podstawie wykresów",1)
    counter = 1 
    document.add_heading("Testowanie funkcji kwantyzującej",2)
    for file in SinFiles:
        Signal, Fs = sf.read(os.path.join(AudioDir,file["File"]), dtype='float32') 
        for bit in PlotSettings["Bits"]:
            kSignal=Kwant(Signal,bit)
            f,axs=plt.subplots(2,1,num=counter,figsize=(5,5)) 

            plotAudio(Signal=kSignal,Fs=Fs,axs=axs,fsize=2**12,TimeMargin=file["TimeMargin"])
            f.suptitle(f"{file['File']} Kwantyzacja {bit}-bitów")
            memfile = BytesIO() 
            f.savefig(memfile)
            document.add_picture(memfile, width=Inches(6)) # set document size
            memfile.close()
            f.clf()
    document.add_heading("Testowanie funkcji decymującej",2)        
    for file in SinFiles:
        Signal, Fs = sf.read(os.path.join(AudioDir,file["File"]), dtype='float32') 
        for step in PlotSettings["Decimation"]:
            dec_Signal,dec_Fs=decimation(Signal,Fs,step)
            f,axs=plt.subplots(2,1,num=counter,figsize=(5,5)) 

            plotAudio(Signal=dec_Signal,Fs=dec_Fs,axs=axs,fsize=2**12,TimeMargin=file["TimeMargin"])
            f.suptitle(f"{file['File']} Decimation step {step}")
            memfile = BytesIO() 
            f.savefig(memfile)
            document.add_picture(memfile, width=Inches(6)) # set document size
            memfile.close()
            f.clf()
    document.add_heading("Testowanie funkcji interpolujących",2)        
    for file in SinFiles:
        Signal, Fs = sf.read(os.path.join(AudioDir,file["File"]), dtype='float32') 
        for Int_Fs in PlotSettings["InterpolationFrequency"]:
            for kind in Interpolation_kind:
                Int_Signal=interpolation(Signal=Signal,Fs=Fs,NewFs=Int_Fs,kind=kind)
                f,axs=plt.subplots(2,1,num=counter,figsize=(5,5)) 

                handle=plotAudio(Signal=Int_Signal,Fs=Int_Fs,axs=axs,fsize=2**12,TimeMargin=file["TimeMargin"])
                f.suptitle(f"{file['File']} Interpolation {kind} Fs {Int_Fs}")
                memfile = BytesIO() 
                f.savefig(memfile)
                document.add_picture(memfile, width=Inches(6)) # set document size
                memfile.close()
                f.clf()
                document.add_paragraph(f"Wykres pomyślnie dodany dla interpolacji: {kind}")
                
    document.add_heading("Podsumowanie pierwszej części zadania",2)
    #document.add_paragraph("")
    document.add_heading("Obserwacje na podstawie odsłuchanych plików ",1)
    
    # generating of audio files
    for file in SingFiles:
        Signal, Fs = sf.read(os.path.join(AudioDir,file), dtype='float32') 
        sfile=file.split(os.sep)[-1].split('.')
        for bit in ListeningSettings["Bits"]:
            kSignal=Kwant(Signal,bit)
            nfile=f"{sfile[0]}_kwant_{bit}.wav"
            sf.write(os.path.join(OutputFolder,nfile),data=kSignal,samplerate=Fs)
        for step in ListeningSettings["Decimation"]:
            dec_Signal,dec_Fs=decimation(Signal,Fs,step)
            
            nfile=f"{sfile[0]}_dec_{step}.wav"
            sf.write(os.path.join(OutputFolder,nfile),data=dec_Signal,samplerate=dec_Fs)
            
        for Int_Fs in ListeningSettings["InterpolationFrequency"]:
            for kind in Interpolation_kind:
                Int_Signal=interpolation(Signal=Signal,Fs=Fs,NewFs=Int_Fs,kind=kind)
                
                nfile=f"{sfile[0]}_interp_{kind}_{Int_Fs}.wav"
                sf.write(os.path.join(OutputFolder,nfile),data=Int_Signal,samplerate=Int_Fs)
        
        
    #document.add_paragraph("")
    document.add_heading("Wnioski ogolne",1)
    document.add_paragraph("Decymacja - dobor odpowiedniego kroku jest kluczowy. Przy korku 24 dla sin8000 trafiamy akurat w takie pubkty, ze sygnal jest (prawie) rowny zero. Przy odsluchach duzy krok powoduje utrate klarownosci sygnalu, pojawiaja sie 'kosmiczne' efekty." \
    "KWANTYZACJA - Szum kwantyzacji - przy zmienjszeniu glebi bitowej na wykresach widoczne bylo zattacenie gladkosci sygnalu i pojawienie sie schodkowania" \
    "KWANTYZACJA - Odsłuch - wyrazne, niemile dla ucha dziweki, prawie jak przestery. Im wiecej bitow tym sygnal bardziej przypomina oryginal." \
    "INTERPOLACJA - obie metody (cubic i linear) radza sobie podobnie, przynajmniej na testach dla sygnalu 60hz. Odsluchowo obie metody przy 4000 tworzyly znane juz z decymacji kosmiczne dzwieki.")
    document.save(OutputRaportFile)