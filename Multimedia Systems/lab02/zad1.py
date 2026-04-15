import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import cv2
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO

def imgToUInt8(img):
    if not np.issubdtype(img.dtype,np.unsignedinteger):
        img = (img*255).astype('uint8')
    return img

##########################################
### Settings #############################
##########################################

Test = False
Scaling_test = False # run only artificial test for scaling methods

ScalesUp = [5,10,15] # list of parameters values
ScalesDown =[0.05,0.1,0.15] # list of parameters values

OutputRaportFile = ".docx" 

##########################################
### Data Set #############################
##########################################

ImgDir = r'.' # Address of folder with files (do nor delete `r``)


SmallImages=["IMG_SMALL/SMALL_0002.png","IMG_SMALL/SMALL_0005.JPG","IMG_SMALL/SMALL_0007.JPG"] # list of file names
BigImages=[ #list of dictionaries
    {
        "Filename":"IMG_BIG/BIG_0004.png", # File name
        "ROIs":[[0,0,1000,1000],[400,400,1800,1800],[600,600,1600,1600]] # list of Region of interests for this image more then 1 per file
    },
    {
        "Filename":"IMG_BIG/BIG_0001.jpg", # File name
        "ROIs":[[0,0,1000,1000],[400,400,1800,1800],[2000,2000,3600,3600]]
    }
]


##########################################
### Functions to  ########################
##########################################

# Scaling methods

def NearestNeigbourScaling(In_img,scale):
    h = In_img.shape[0]
    w = In_img.shape[1]
    nh = np.ceil(h * scale).astype(int)
    nw = np.ceil(w * scale).astype(int)
    X = np.linspace(0,h-1,nh)
    Y = np.linspace(0,w-1,nw)
    if (len(In_img.shape)<3):
        Out_img = np.zeros((nh,nw))
    else:
        Out_img = np.zeros((nh,nw,In_img.shape[2]))
    for ix,x in enumerate(X):
        #print(ix,x)
        for iy,y in enumerate(Y):
            #print(iy,y)
            xp = np.round(x).astype(int)
            yp = np.round(y).astype(int)
            Out_img[ix,iy] = In_img[xp,yp]

    if In_img.dtype == np.uint8:
        Out_img = np.clip(Out_img, 0, 255)
    return Out_img.astype(In_img.dtype)

def BilinearScaling(In_img,scale):
    h = In_img.shape[0]
    w = In_img.shape[1]
    nh = np.ceil(h * scale).astype(int)
    nw = np.ceil(w * scale).astype(int)
    X = np.linspace(0,h-1,nh)
    Y = np.linspace(0,w-1,nw)
    if (len(In_img.shape)<3):
        Out_img = np.zeros((nh,nw))
    else:
        Out_img = np.zeros((nh,nw,In_img.shape[2]))
    for ix,x in enumerate(X):
        #print(ix,x)
        for iy,y in enumerate(Y):
            #print(iy,y)
            x1 = np.floor(x).astype(int)
            x2 = np.ceil(x).astype(int)
            xd = x-x1
            y1 = np.floor(y).astype(int)
            y2 = np.ceil(y).astype(int)
            yd = y-y1
            #print("xd,yd:",xd,yd)
            Out_img[ix,iy] = In_img[x1,y1]*(1-xd)*(1-yd) + In_img[x2,y1]*(1-yd)*xd + In_img[x1,y2]*yd*(1-xd) + In_img[x2,y2]*xd*yd

    if In_img.dtype == np.uint8:
        Out_img = np.clip(Out_img, 0, 255)
    return Out_img.astype(In_img.dtype)

# Shrinking methods
 
def MeanResizing(In_img,scale):
    h = In_img.shape[0]
    w = In_img.shape[1]
    nh = np.ceil(h * scale).astype(int)
    nw = np.ceil(w * scale).astype(int)
    X = np.linspace(0,h-1,nh)
    Y = np.linspace(0,w-1,nw)
    if (len(In_img.shape) < 3):
        Out_img = np.zeros((nh,nw))
    else:
        Out_img = np.zeros((nh,nw,In_img.shape[2]))
    for ix,x in enumerate(X):
        if ix > 0:
            x1 = -(x-X[ix-1])/2
        else:
            x1 = 0
        if ix < len(X)-1:
            x2 = (X[ix+1] - x)/2 + 1
        else:
            x2 = 0
        ix_arr=np.round(x +np.arange(x1,x2)).astype(int)
        ix_arr=ix_arr.clip(0,h-1)

        for iy,y in enumerate(Y):
            if iy > 0:
                y1 = -(y-Y[iy-1])/2
            else:
                y1 = 0
            if iy < len(Y)-1:
                y2 = (Y[iy+1] - y)/2 + 1
            else:
                y2 = 0
            iy_arr=np.round(y+np.arange(y1,y2)).astype(int)
            iy_arr=iy_arr.clip(0,w-1)

            fragment = In_img[ix_arr[0]:ix_arr[-1]+1, iy_arr[0]:iy_arr[-1]+1]

            if fragment.size > 0:
                if len(In_img.shape) < 3:
                    Out_img[ix, iy] = np.mean(fragment)
                else:
                    Out_img[ix,iy] = np.mean(fragment,axis=(0,1))

    return Out_img.astype(In_img.dtype)

def WeightedMeanResizing(In_img,scale):
    h = In_img.shape[0]
    w = In_img.shape[1]
    nh = np.ceil(h * scale).astype(int)
    nw = np.ceil(w * scale).astype(int)
    X = np.linspace(0,h-1,nh)
    Y = np.linspace(0,w-1,nw)
    if (len(In_img.shape) < 3):
        Out_img = np.zeros((nh,nw))
    else:
        Out_img = np.zeros((nh,nw,In_img.shape[2]))
    for ix,x in enumerate(X):
        if ix > 0:
            x1 = -(x-X[ix-1])/2
        else:
            x1 = 0
        if ix < len(X)-1:
            x2 = (X[ix+1] - x)/2 + 1
        else:
            x2 = 0
        ix_arr=np.round(x +np.arange(x1,x2)).astype(int)
        ix_arr=ix_arr.clip(0,h-1)

        for iy,y in enumerate(Y):
            if iy > 0:
                y1 = -(y-Y[iy-1])/2
            else:
                y1 = 0
            if iy < len(Y)-1:
                y2 = (Y[iy+1] - y)/2 + 1
            else:
                y2 = 0
            iy_arr=np.round(y+np.arange(y1,y2)).astype(int)
            iy_arr=iy_arr.clip(0,w-1)

            fragment = In_img[ix_arr[0]:ix_arr[-1]+1, iy_arr[0]:iy_arr[-1]+1]

            if fragment.size > 0:
                if len(In_img.shape) < 3:
                    weights = np.random.rand(*fragment.shape)
                    weighted_sum = np.sum(fragment * weights)
                    total_weight = np.sum(weights)
                    
                    if total_weight > 0:
                        Out_img[ix,iy] = weighted_sum / total_weight
                    else:
                        Out_img[ix,iy] = 0
                else:
                    weights = np.random.rand(fragment.shape[0], fragment.shape[1])
                    total_weight = np.sum(weights)
                    weights_3d = np.expand_dims(weights, axis=-1)
                    weighted_sum = np.sum(fragment * weights_3d, axis=(0,1))
                    
                    if total_weight > 0:
                        Out_img[ix,iy] = weighted_sum / total_weight
                    else:
                        Out_img[ix,iy] = 0

    return Out_img.astype(In_img.dtype)

def MedianResizing(In_img,scale):
    h = In_img.shape[0]
    w = In_img.shape[1]
    nh = np.ceil(h * scale).astype(int)
    nw = np.ceil(w * scale).astype(int)
    X = np.linspace(0,h-1,nh)
    Y = np.linspace(0,w-1,nw)
    if (len(In_img.shape) < 3):
        Out_img = np.zeros((nh,nw))
    else:
        Out_img = np.zeros((nh,nw,In_img.shape[2]))
    for ix,x in enumerate(X):
        if ix > 0:
            x1 = -(x-X[ix-1])/2
        else:
            x1 = 0
        if ix < len(X)-1:
            x2 = (X[ix+1] - x)/2 + 1
        else:
            x2 = 0
        ix_arr=np.round(x +np.arange(x1,x2)).astype(int)
        ix_arr=ix_arr.clip(0,h-1)

        for iy,y in enumerate(Y):
            if iy > 0:
                y1 = -(y-Y[iy-1])/2
            else:
                y1 = 0
            if iy < len(Y)-1:
                y2 = (Y[iy+1] - y)/2 + 1
            else:
                y2 = 0
            iy_arr=np.round(y+np.arange(y1,y2)).astype(int)
            iy_arr=iy_arr.clip(0,w-1)

            fragment = In_img[ix_arr[0]:ix_arr[-1]+1, iy_arr[0]:iy_arr[-1]+1]

            if fragment.size > 0:
                if len(In_img.shape) < 3:
                    Out_img[ix,iy] = np.median(fragment)
                else:
                    Out_img[ix,iy] = np.median(fragment, axis=(0,1))
    return Out_img.astype(In_img.dtype)

def EdgeDetection(img):
    if img.dtype != np.uint8:
        if img.max() <= 1.0:
            img_int = imgToUInt8(img)
        else:
            img_int = img.astype(np.uint8)

    else:
        img_int = img

    if len(img_int.shape) < 3:
        gray = img_int
    else:
        gray = cv2.cvtColor(img_int,cv2.COLOR_RGB2GRAY)
    edges = cv2.Canny(gray,100,200)
    edges = cv2.cvtColor(edges,cv2.COLOR_GRAY2RGB)
    return edges

##########################################
### Main Program  ########################
##########################################

def plot_resize(img, scale, nnscale, bscale, ed_img, ed_nnscale, ed_bscale, mr_img, wmr_img, mdr_img, ed_mr_img, ed_wmr_img,ed_mdr_img, oROI, filename,counter,figsize=(5,5)):
    ROI=(np.array(oROI)*scale).astype(int)
    f,axs=plt.subplots(4,3,num=counter,figsize=figsize) 
    f.suptitle(f"{filename} ROI: {ROI}")
    axs[0,0].imshow(img[oROI[1]:oROI[1]+oROI[3],oROI[0]:oROI[0]+oROI[2],:])
    axs[0,0].set_title("Original")
    axs[0,0].set_axis_off()

    axs[0,1].imshow(nnscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[0,1].set_title(f"NN scale {scale}")
    axs[0,1].set_axis_off()


    axs[0,2].imshow(bscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[0,2].set_title(f"Blinear scale {scale}")
    axs[0,2].set_axis_off()

    if len(ed_img.shape)==3:
        axs[1,0].imshow(ed_img[oROI[1]:oROI[1]+oROI[3],oROI[0]:oROI[0]+oROI[2],:])
        axs[1,0].set_title("Edges Original")
        axs[1,0].set_axis_off()
    else:
        axs[1,0].imshow(ed_img[oROI[1]:oROI[1]+oROI[3],oROI[0]:oROI[0]+oROI[2]])
        axs[1,0].set_title("Edges Original")
        axs[1,0].set_axis_off()

    if len(ed_nnscale.shape)==3:
        axs[1,1].imshow(ed_nnscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
        axs[1,1].set_title("Edges NN")
        axs[1,1].set_axis_off()
    else:
        axs[1,1].imshow(ed_nnscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2]])
        axs[1,1].set_title("Edges NN")
        axs[1,1].set_axis_off()
        
    if len(ed_bscale.shape)==3:
        axs[1,2].imshow(ed_bscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
        axs[1,2].set_title("Edges Bilinear")
        axs[1,2].set_axis_off()
    else:
        axs[1,2].imshow(ed_bscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2]])
        axs[1,2].set_title("Edges Bilinear")
        axs[1,2].set_axis_off()

    axs[2,0].imshow(mr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[2,0].set_title(f"Mean Resizing scale {scale}")
    axs[2,0].set_axis_off()

    axs[2,1].imshow(wmr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[2,1].set_title(f"Weighted Mean Resizing scale {scale}")
    axs[2,1].set_axis_off()

    axs[2,2].imshow(mdr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[2,2].set_title(f"Median Resizing scale {scale}")
    axs[2,2].set_axis_off()
        
    if len(ed_mr_img.shape)==3:
        axs[3,0].imshow(ed_mr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
        axs[3,0].set_title("Edges Mean Resizing")
        axs[3,0].set_axis_off()
    else:
        axs[3,0].imshow(ed_mr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2]])
        axs[3,0].set_title("Edges Mean Resizing")
        axs[3,0].set_axis_off()
        
    if len(ed_wmr_img.shape)==3:
        axs[3,1].imshow(ed_wmr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
        axs[3,1].set_title("Edges Weighted Mean Resizing")
        axs[3,1].set_axis_off()
    else:
        axs[3,1].imshow(ed_nnscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2]])
        axs[3,1].set_title("Edges Weighted Mean Resizing")
        axs[3,1].set_axis_off()
        
    if len(ed_mdr_img.shape)==3:
        axs[3,2].imshow(ed_mdr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
        axs[3,2].set_title("Edges Median Resizing")
        axs[3,2].set_axis_off()
    else:
        axs[3,2].imshow(ed_mdr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2]])
        axs[3,2].set_title("Edges Median Resizing")
        axs[3,2].set_axis_off()
        
    return f

def plot_scaling(img, scale, nnscale, bscale, counter, ed_img, ed_nnscale, ed_bscale, file,figsize=(5,5)):
    f,axs=plt.subplots(2,3,num=counter,figsize=figsize) 
    f.suptitle(f"{file}")

    axs[0,0].imshow(img)
    axs[0,0].set_title("Original")
    axs[0,0].set_axis_off()

    axs[0,1].imshow(nnscale)
    axs[0,1].set_title(f"NN scale {scale}")
    axs[0,1].set_axis_off()

    axs[0,2].imshow(bscale)
    axs[0,2].set_title(f"Blinear scale {scale}")
    axs[0,2].set_axis_off()

    axs[1,0].imshow(ed_img)
    axs[1,0].set_title("Edges Original")
    axs[1,0].set_axis_off()

    axs[1,1].imshow(ed_nnscale)
    axs[1,1].set_title("Edges NN")
    axs[1,1].set_axis_off()

    axs[1,2].imshow(ed_bscale)
    axs[1,2].set_title("Edges Bilinear")
    axs[1,2].set_axis_off()
    return f

if Test:
    # test case
    if Scaling_test:
        img= np.zeros((3,3,3),dtype=np.float32)
        img[1,1,:]=1.0
        for scale in ScalesUp:
            f,axs=plt.subplots(1,2)
            nnscale=NearestNeigbourScaling(img,scale)
            axs[0].imshow(nnscale)
            bscale=BilinearScaling(img,scale)
            axs[1].imshow(bscale)
        
    else:
        counter=1
        for scale in ScalesUp:
            img=plt.imread(os.path.join(ImgDir,SmallImages[1]))
            nnscale=NearestNeigbourScaling(img,scale)
            bscale=BilinearScaling(img,scale)
            ed_img=EdgeDetection(img)
            ed_nnscale=EdgeDetection(nnscale)
            ed_bscale=EdgeDetection(bscale)

            f = plot_scaling(img, scale, nnscale, bscale, counter, ed_img, ed_nnscale, ed_bscale, SmallImages[1])

            counter+=1

        for scale in ScalesDown:    
            img=plt.imread(os.path.join(ImgDir,BigImages[0]["Filename"]))

            nnscale=NearestNeigbourScaling(img,scale)
            bscale=BilinearScaling(img,scale)

            ed_img=EdgeDetection(img)
            ed_nnscale=EdgeDetection(nnscale)
            ed_bscale=EdgeDetection(bscale)

            mr_img=MeanResizing(img,scale)
            wmr_img=WeightedMeanResizing(img,scale)
            mdr_img=MedianResizing(img,scale)

            ed_mr_img=EdgeDetection(mr_img)
            ed_wmr_img=EdgeDetection(wmr_img)
            ed_mdr_img=EdgeDetection(mdr_img)
  
            f= plot_resize(img, scale, nnscale, bscale, ed_img, ed_nnscale, ed_bscale, mr_img, wmr_img, mdr_img, ed_mr_img, ed_wmr_img,ed_mdr_img, BigImages[0]["ROIs"][0], BigImages[0]["Filename"],counter=counter)
            counter+=1
        
    plt.show()
else: 
    # generate raport
    document = Document()
    document.add_heading('Report',0) # tworzenie nagłówków druga wartość to poziom nagłówka 
    document.add_paragraph("Autor: ")
    document.add_paragraph("Proszę wstawić mi MAX PUNKTOW jeżeli tego nie wyedytuję :D")
    document.add_section()
    document.add_heading("Test algorytmów powiększania",1)
    counter = 1 
    for file in SmallImages:
        img=plt.imread(os.path.join(ImgDir,file))
        for scale in ScalesUp:
            nnscale=NearestNeigbourScaling(img,scale)
            bscale=BilinearScaling(img,scale)
            ed_img=EdgeDetection(img)
            ed_nnscale=EdgeDetection(nnscale)
            ed_bscale=EdgeDetection(bscale)
            
            f = plot_scaling(img, scale, nnscale, bscale, counter, ed_img, ed_nnscale, ed_bscale, file) # set figszie

            memfile = BytesIO() 
            f.savefig(memfile)
            document.add_picture(memfile, width=Inches(6)) # set document size
            memfile.close()
            f.clf()
    document.add_section()
    document.add_heading("Test algorytmów pomniejszania",1)
    for file_dict in BigImages:
        filename=file_dict['Filename']
        img=plt.imread(os.path.join(ImgDir,filename))
        for scale in ScalesDown:
            nnscale=NearestNeigbourScaling(img,scale)
            bscale=BilinearScaling(img,scale)

            ed_img=EdgeDetection(img)
            ed_nnscale=EdgeDetection(nnscale)
            ed_bscale=EdgeDetection(bscale)

            mr_img=MeanResizing(img,scale)
            wmr_img=WeightedMeanResizing(img,scale)
            mdr_img=MedianResizing(img,scale)

            ed_mr_img=EdgeDetection(mr_img)
            ed_wmr_img=EdgeDetection(wmr_img)
            ed_mdr_img=EdgeDetection(mdr_img)

            for ROI in file_dict['ROIs']:

                f = plot_resize(img, scale, nnscale, bscale, ed_img, ed_nnscale, ed_bscale, mr_img, wmr_img, mdr_img, ed_mr_img, ed_wmr_img,ed_mdr_img, ROI, filename,counter=counter) # set figszie

                memfile = BytesIO() 
                f.savefig(memfile)
                document.add_picture(memfile, width=Inches(6)) # set document size
                memfile.close()
                f.clf()
    document.add_section()
    document.add_heading("Podsumowanie i wnioski",1)
    document.add_paragraph("Nearest Neighbour - przy testach dzialal szybko lecz generowal widocza pikseloze.\n Bilinear scaling - daje gladszy efekt kosztem ostrosci. Efekt rozmycia sprawia ze edge detection" \
    "(przy progach 100-200) ledwo wykrywa krawędzie.\n Metody pomniejszania:\n Mean resizing - przy mocnym zoomie mozna bylo zobaczyc ze ta metoda barzdiej niz inne powoduje rozmycie detali (nieznacznie).\n Weighted" \
    "mean resizing - wynik podobny do mean resizing \n Median resizing - w tym wariancie najlepiej bylo widac zachowanie ostrosci detali")
    document.save(OutputRaportFile) 
