import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import soundfile as sf
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO

##########################################
### Settings #############################
##########################################

Only_Tests=True
bit_test=8

DPCM_n=5
DPCM_predictor=np.mean

OutputRaportFile = "raport.docx" 
OutputFolder="audio" # place for all your new audio files will be

##########################################
### Data Set #############################
##########################################

AudioDir = r'.' # Address of folder with files (do nor delete `r``)

SingFiles=['SING/sing_high1.wav','SING/sing_low1.wav','SING/sing_medium1.wav'] # list of file names of with Singing Voice

##########################################
### Functions to  ########################
##########################################

def Kwant(x,bit):
    d = (2 ** bit) - 1 
    
    if np.issubdtype(x.dtype, np.floating):
        d_min = -1
        d_max = 1
    else:
        d_min = np.iinfo(x.dtype).min
        d_max = np.iinfo(x.dtype).max

    DataF = x.astype(float)

    if d_max > d_min:
        DataF = (DataF - d_min) / (d_max - d_min)
        DataF = np.round(DataF * d) / d
        DataF = DataF * (d_max - d_min) + d_min
    return DataF.astype(x.dtype)

def A_law_compress(x):
    y = np.copy(x)
    A = 87.6
    x_abs = abs(y)
    idx = (x_abs < (1/A))
    s = np.sign(y)
    y[idx] = s[idx] * (A * x_abs[idx]) / (1 + np.log(A))
    y[np.logical_not(idx)] = s[np.logical_not(idx)] * (1+np.log(A*x_abs[np.logical_not(idx)]))/(1 + np.log(A))
    return y

def A_law_decompress(x):
    y = np.copy(x)
    A = 87.6
    y_abs = abs(y)
    s = np.sign(y)
    idx = (y_abs < (1/A))
    y[idx] = s[idx] * (y_abs[idx] * (1+ np.log(A))) / A
    y[np.logical_not(idx)] = s[np.logical_not(idx)] * (np.exp(y_abs[np.logical_not(idx)] * (1 + np.log(A)) - 1)) / A
    return y
    
def mu_law_compress(x):
    y=x
    mu=255
    x_abs = abs(y)
    s = np.sign(y)
    idx = (-1 <= y) & (y <= 1)
    y = s[idx] * (np.log(1+mu*x_abs[idx]))/np.log(1+mu)
    return y

def mu_law_decompress(x):
    y=x
    mu = 255
    s = np.sign(y)
    y_abs = abs(y)
    idx = (-1 <= y) & (y <= 1)
    y = s[idx] * (1/mu) * (((1+mu)**y_abs) - 1)
    return y

def DPCM_compress(x,bit):
    y=np.zeros(x.shape)
    e=0
    for i in range(0,x.shape[0]):
        y[i]=Kwant(x[i]-e,bit)
        e+=y[i]
    return y

def DPCM_decompress(x):
    y = np.zeros(x.shape)
    e = 0
    for i in range(0, x.shape[0]):
        y[i] = x[i] + e
        e += x[i]
    return y

def DPCM_compress_pred(x,bit,n,predictor=np.mean): 
    y=np.zeros(x.shape)
    xp=np.zeros(x.shape)
    e=0
    for i in range(0,x.shape[0]):
        y[i]=Kwant(x[i]-e,bit)
        xp[i]=y[i]+e
        idx=(np.arange(i-n,i,1,dtype=int)+1)
        idx=np.delete(idx,idx<0)
        e=predictor(xp[idx])
    return y

def DPCM_decompress_pred(x,n,predictor=np.mean):
    y = np.zeros(x.shape)
    e = 0
    for i in range(0, x.shape[0]):
        y[i] = x[i] + e
        
        idx = (np.arange(i - n, i, 1, dtype=int) + 1)
        idx = np.delete(idx, idx < 0)
        e = predictor(y[idx])
    return y

##########################################
### Main Program  ########################
##########################################


document = Document()
if not Only_Tests:
    # generate raport
    document.add_heading('Report',0) # tworzenie nagłówków druga wartość to poziom nagłówka 
    document.add_paragraph("Autor: Dorian Sobierański")
    document.add_section()
    document.add_heading('Wykresy testujące działanie algorytmów',1)

x=np.linspace(-1,1,1000)
y=0.9*np.sin(np.pi*x*4)

x_alaw_comp=A_law_compress(x)
x_alaw_comp=Kwant(x_alaw_comp,bit_test)
x_alaw_decomp=A_law_decompress(x_alaw_comp)

x_mulaw_comp=mu_law_compress(x)
x_mulaw_comp=Kwant(x_mulaw_comp,bit_test)
x_mulaw_decomp=mu_law_decompress(x_mulaw_comp)

y_alaw_decomp =A_law_decompress(Kwant(A_law_compress(y),bit_test))
y_mulaw_decomp =mu_law_decompress(Kwant(mu_law_compress(y),bit_test))

dpcm_c=DPCM_compress(y,bit_test)
dpcm_dec=DPCM_decompress(dpcm_c)
dpcm_c_p=DPCM_compress_pred(y,bit_test,n=DPCM_n,predictor=DPCM_predictor)
dpcm_dec_p=DPCM_decompress_pred(dpcm_c_p,n=DPCM_n,predictor=DPCM_predictor)

f1,axs=plt.subplots(1,2,num=1,figsize=(6,6)) 
f1.suptitle(f"Test kompresji law dla {bit_test} bitów")
axs[0].plot(x,x_alaw_comp,label="a_law")
axs[0].plot(x,x_mulaw_comp,label="mu_law")
axs[0].set_title("Sygnał po kompresji")
axs[0].legend()

axs[1].plot(x,x_alaw_decomp,label="a_law")
axs[1].plot(x,x_mulaw_decomp,label="mu_law")
axs[1].set_title("Sygnał po dekompresji")
axs[1].legend()

f2,axs=plt.subplots(5,1,num=2,figsize=(8,6)) 
f2.suptitle(f"Test dekompresji dla {bit_test} bitów")
axs[0].plot(x,y,label="Sygnał bazowy")
axs[0].set_title("Sygnał bazowy")

axs[1].plot(x,y_alaw_decomp,label="Sygnał po kompresji A-law")
axs[1].legend()

axs[2].plot(x,y_mulaw_decomp,label="Sygnał po kompresji mu-law")
axs[2].legend()

axs[3].plot(x,dpcm_dec,label="Sygnał po kompresji DPCM bez predykcji")
axs[3].legend()

axs[4].plot(x,dpcm_dec_p,label="Sygnał po kompresji DPCM z predykcją")
axs[4].legend()


if Only_Tests:
    plt.show()
else:
    memfile = BytesIO() 
    f1.savefig(memfile)
    document.add_picture(memfile, width=Inches(6)) # set document size
    memfile.close()
    f1.clf()
    memfile = BytesIO() 
    f2.savefig(memfile)
    document.add_picture(memfile, width=Inches(6)) # set document size
    memfile.close()
    f2.clf()  
    document.add_section()
    document.add_heading("Obserwacje na podstawie odsłuchanych plików ",1)
    document.add_paragraph("Z wykresów i odsłuchu wynika, że metody a-law i mu-law kompresują głównie głośne dźwięki, "
        "zachowując dużą precyzję dla cichych detali. Dlatego wokale brzmią po nich bardzo naturalnie. "
        "Z kolei DPCM nie kompresuje samej głośności, lecz różnice między kolejnymi próbkami. "
        "W praktyce algorytm ten świetnie radzi sobie ze spokojnym śpiewem, ale gubi się i zauważalnie zniekształca dźwięk "
        "przy nagłych, ostrych skokach częstotliwości.")
    document.add_heading("Zadanie 2.2",2)
    document.add_paragraph("Dla standardowej wartości 8 bitów, jakość dźwięku we wszystkich metodach jest bardzo dobra. "
        "Pliki po kompresji logarytmicznej brzmią naturalnie, słychać jedynie lekki szum "
        "oraz taki telefoniczny efekt. Sygnał odtworzony za pomocą metody DPCM również jest "
        "w pełni zrozumiały, przy czym slychać mocniej niż w innych metodach szum. Nie słyszałem wyraźnej różnicy między predykcją a jej brakiem.")
    document.add_heading("Zadanie 2.3",2)
    document.add_paragraph("Z przeprowadzonych testów wynika, że przy zejściu na 7 i 6 bitów sygnał pozostaje w pełni zrozumiały, "
        "choć stopniowo narasta szum kwantyzacji. Przy 5 i 4 bitach jakość drastycznie spada, "
        "ale treść wokalu wciąż można z trudem odszyfrować. Poniżej 4 bitów dźwięk staje się całkowicie "
        "niezrozumiałym trzeszczeniem. Dodatkowo - przy sing_high1_mu_LAW_3b w połowie dźwięku następuje ogromne podgłoszenie, które nie jest najmilszym doświadczeniem" \
        "gdy ma się słuchawki. DCPM powoduje występowanie szumów już przy 7 bitach, podczas gdy mu i a law trzyma się wtedy dobrze.")
    document.add_section()
    document.add_heading("Podsumowanie i Wnioski",1)
    document.add_paragraph(
        "1. Metody a-law, mu-law maskują szum dla cichych dźwięków, co idealnie pokrywa się z właściwościami ludzkiego słuchu.\n"
        "2. Zastosowanie predykcji w algorytmie DPCM jest kluczowe, ponieważ poprawia jakość dźwięku i stabilność względem nagłych zmian sygnału.\n"
        "3. Redukcja informacji poniżej 4-5 bitów całkowicie niszczy jakość.\n"
        "4. Wokale 'high' szybciej tracą jakość, ponieważ gwałtowne zmiany w wysokich rejestrach są trudniejsze do skompresowania i przewidzenia algorytmem DPCM." \
        "Analizowane pliki - sing_high1, sing_low1, sing_medium1"
    )
    document.save(OutputRaportFile) 
    if OutputFolder != "":
        os.makedirs(OutputFolder, exist_ok=True)
    for file in SingFiles:
        Signal, Fs = sf.read(os.path.join(AudioDir,file), dtype='float32') 
        
        nazwa_pliku = os.path.basename(file)
        sfile_name = os.path.splitext(nazwa_pliku)[0]
        
        for bit in [8,7,6,5,4,3,2]:
            y_alaw_decomp = A_law_decompress(Kwant(A_law_compress(Signal), bit))
            y_mulaw_decomp = mu_law_decompress(Kwant(mu_law_compress(Signal), bit))

            dpcm_c = DPCM_compress(Signal, bit)
            dpcm_dec = DPCM_decompress(dpcm_c)
            dpcm_c_p = DPCM_compress_pred(Signal, bit, n=DPCM_n, predictor=DPCM_predictor)
            dpcm_dec_p = DPCM_decompress_pred(dpcm_c_p, n=DPCM_n, predictor=DPCM_predictor)
            
            sf.write(os.path.join(OutputFolder, f"{sfile_name}_A_LAW_{bit}b.wav"), data=y_alaw_decomp, samplerate=Fs)
            sf.write(os.path.join(OutputFolder, f"{sfile_name}_mu_LAW_{bit}b.wav"), data=y_mulaw_decomp, samplerate=Fs)
            sf.write(os.path.join(OutputFolder, f"{sfile_name}_DPCM_bp_{bit}b.wav"), data=dpcm_dec, samplerate=Fs)
            sf.write(os.path.join(OutputFolder, f"{sfile_name}_DPCM_zp_{bit}b.wav"), data=dpcm_dec_p, samplerate=Fs)

            